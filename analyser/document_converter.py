# analyser/document_converter.py

import logging
import io
from typing import Tuple, Optional
from docx import Document as DocxDocument
from openpyxl import load_workbook
from PyPDF2 import PdfReader

logger = logging.getLogger(__name__)


class DocumentConverter:
    """Conversor de documentos para extrair conteúdo em diferentes formatos."""
    
    # Formatos suportados
    FORMATOS_SUPORTADOS = {
        'application/pdf': 'PDF',
        'application/vnd.openxmlformats-officedocument.wordprocessingml.document': 'DOCX',
        'application/msword': 'DOC',
        'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet': 'XLSX',
        'application/vnd.ms-excel': 'XLS',
        'text/plain': 'TXT',
    }
    
    @staticmethod
    def is_supported(mime_type: str) -> bool:
        """Verifica se o MIME type é suportado."""
        return mime_type in DocumentConverter.FORMATOS_SUPORTADOS
    
    @staticmethod
    def get_format_type(mime_type: str) -> Optional[str]:
        """Retorna o tipo de formato a partir do MIME type."""
        return DocumentConverter.FORMATOS_SUPORTADOS.get(mime_type)
    
    @staticmethod
    def extract_text_from_docx(content: bytes) -> str:
        """Extrai texto de um arquivo DOCX."""
        logger.info("📄 Extraindo texto de DOCX...")
        try:
            doc = DocxDocument(io.BytesIO(content))
            texto = []
            
            for para in doc.paragraphs:
                if para.text.strip():
                    texto.append(para.text)
            
            # Extrai também de tabelas
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells]
                    if any(row_text):
                        texto.append(' | '.join(row_text))
            
            resultado = '\n'.join(texto)
            logger.info(f"✅ Extraído {len(resultado)} caracteres do DOCX")
            return resultado
        
        except Exception as e:
            logger.error(f"❌ Erro ao extrair DOCX: {e}")
            raise
    
    @staticmethod
    def extract_text_from_xlsx(content: bytes) -> str:
        """Extrai dados de um arquivo XLSX."""
        logger.info("📊 Extraindo dados de XLSX...")
        try:
            wb = load_workbook(io.BytesIO(content))
            texto = []
            
            for sheet_name in wb.sheetnames:
                sheet = wb[sheet_name]
                texto.append(f"\n=== Planilha: {sheet_name} ===\n")
                
                for row in sheet.iter_rows(values_only=True):
                    row_text = [str(cell) if cell is not None else '' for cell in row]
                    if any(row_text):
                        texto.append(' | '.join(row_text))
            
            resultado = '\n'.join(texto)
            logger.info(f"✅ Extraído {len(resultado)} caracteres do XLSX")
            return resultado
        
        except Exception as e:
            logger.error(f"❌ Erro ao extrair XLSX: {e}")
            raise
    
    @staticmethod
    def extract_text_from_pdf(content: bytes) -> str:
        """Extrai texto de um arquivo PDF."""
        logger.info("📕 Extraindo texto de PDF...")
        try:
            pdf_reader = PdfReader(io.BytesIO(content))
            texto = []
            
            for page_num, page in enumerate(pdf_reader.pages, 1):
                page_text = page.extract_text()
                if page_text.strip():
                    texto.append(f"--- Página {page_num} ---\n{page_text}")
            
            resultado = '\n'.join(texto)
            logger.info(f"✅ Extraído {len(resultado)} caracteres do PDF")
            return resultado
        
        except Exception as e:
            logger.error(f"❌ Erro ao extrair PDF: {e}")
            raise
    
    @staticmethod
    def extract_text_from_doc(content: bytes) -> str:
        """Extrai texto de um arquivo DOC (tentando como DOCX primeiro)."""
        logger.info("📄 Extraindo texto de DOC...")
        try:
            # Tenta primeiro como DOCX (muitos arquivos .doc modernos são DOCX)
            return DocumentConverter.extract_text_from_docx(content)
        except Exception as e:
            logger.warning(f"⚠️ Não foi possível extrair como DOCX, erro: {e}")
            # Se falhar, tenta extrair como texto simples
            try:
                return content.decode('utf-8', errors='ignore')
            except Exception as e2:
                logger.error(f"❌ Erro ao extrair DOC: {e2}")
                raise
    
    @staticmethod
    def convert_to_text(file_content: bytes, mime_type: str, file_name: str) -> Tuple[str, str]:
        """
        Converte documento para texto.
        
        :param file_content: Conteúdo binário do arquivo
        :param mime_type: MIME type do arquivo
        :param file_name: Nome do arquivo (para logging)
        :return: Tuple (texto_extraído, formato)
        """
        
        logger.info(f"🔄 Convertendo arquivo: {file_name} (MIME: {mime_type})")
        
        if not DocumentConverter.is_supported(mime_type):
            raise ValueError(f"❌ Formato não suportado: {mime_type}")
        
        formato = DocumentConverter.get_format_type(mime_type)
        
        # Roteia para o método apropriado
        if formato == 'PDF':
            texto = DocumentConverter.extract_text_from_pdf(file_content)
        elif formato == 'DOCX':
            texto = DocumentConverter.extract_text_from_docx(file_content)
        elif formato == 'DOC':
            texto = DocumentConverter.extract_text_from_doc(file_content)
        elif formato == 'XLSX':
            texto = DocumentConverter.extract_text_from_xlsx(file_content)
        elif formato == 'XLS':
            # XLS antigo - tenta como XLSX
            texto = DocumentConverter.extract_text_from_xlsx(file_content)
        elif formato == 'TXT':
            texto = file_content.decode('utf-8', errors='ignore')
        else:
            raise ValueError(f"❌ Formato não implementado: {formato}")
        
        logger.info(f"✅ Conversão concluída: {formato} -> Texto ({len(texto)} caracteres)")
        return texto, formato


# Exemplo de uso:
if __name__ == '__main__':
    # Para testar
    with open('exemplo.docx', 'rb') as f:
        content = f.read()
        texto, formato = DocumentConverter.convert_to_text(
            content,
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            'exemplo.docx'
        )
        print(f"Formato: {formato}")
        print(f"Texto extraído:\n{texto[:500]}...")